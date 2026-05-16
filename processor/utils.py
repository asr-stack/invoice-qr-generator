import pandas as pd
from docx import Document
import re
import io
import qrcode
import os
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO

def read_excel(file):
    df = pd.read_excel(file)

    df.columns = df.columns.str.strip().str.lower()

    print("Columns found:", df.columns)

    doc_col = None
    irn_col = None
    ack_no_col = None
    ack_date_col = None

    for col in df.columns:
        if "doc no" in col:
            doc_col = col
        elif col.strip() == "irn":
            irn_col = col
        elif "ack no" in col:
            ack_no_col = col
        elif "ack date" in col:
            ack_date_col = col

    if not all([doc_col, irn_col, ack_no_col, ack_date_col]):
        raise Exception("Missing required columns")

    print("Using columns:", doc_col, irn_col, ack_no_col, ack_date_col)

    df[doc_col] = df[doc_col].astype(str).str.strip()

    mapping = {}

    for _, row in df.iterrows():
        mapping[row[doc_col]] = {
            "irn": str(row[irn_col]).strip(),
            "ack_no": str(row[ack_no_col]).strip(),
            "ack_date": str(row[ack_date_col]).strip()
        }

    print("Mapping:", mapping)

    return mapping

def extract_invoice_number(file):
    doc = Document(file)

    full_text = "\n".join([p.text for p in doc.paragraphs])

    # Include tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    print("🔍 FULL TEXT:\n", full_text)

    # Look for long invoice-like numbers
    matches = re.findall(r'\b\d{10,15}\b', full_text)

    if matches:
        invoice_no = matches[0]
        print("✅ Extracted invoice:", invoice_no)
        return invoice_no

    print("❌ Invoice not found")
    return None

def generate_qr(irn):
    img = qrcode.make(irn)

    qr_bytes = BytesIO()
    img.save(qr_bytes, format='PNG')

    qr_bytes.seek(0)

    return qr_bytes



def insert_qr_into_docx(input_docx_path, qr_image, irn, ack_no, ack_date, original_name):
    doc = Document(input_docx_path)

    inserted = False

    for para in doc.paragraphs:
        if "Customer VAT ID" in para.text:

            # 👉 LEFT ALIGN
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 👉 Add QR
            run = para.add_run()
            run.add_break()
            run.add_picture(qr_image, width=Inches(1.6))

            # 👉 Add spacing line
            run = para.add_run()
            run.add_break()

            # 👉 FORMATTED TEXT WITH SEPARATORS
            r1 = para.add_run("IRN: ")
            r1.bold = True

            r2 = para.add_run(f"{irn}   |   ")
            r2.bold = False

            r3 = para.add_run("Ack No: ")
            r3.bold = True

            r4 = para.add_run(f"{ack_no}   |   ")
            r4.bold = False

            r5 = para.add_run("Date: ")
            r5.bold = True

            r6 = para.add_run(f"{ack_date}")
            r6.bold = False

            # 👉 APPLY FONT TO ALL
            for r in [r1, r2, r3, r4, r5, r6]:
                r.font.name = "Arial"
                r.font.size = Pt(7)

            inserted = True
            break

    # 👉 fallback
    if not inserted:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = para.add_run()
        run.add_picture(qr_image, width=Inches(1.6))
        run.add_break()

        r1 = para.add_run("IRN: ")
        r1.bold = True

        r2 = para.add_run(f"{irn}   |   ")
        r2.bold = False

        r3 = para.add_run("Ack No: ")
        r3.bold = True

        r4 = para.add_run(f"{ack_no}   |   ")
        r4.bold = False

        r5 = para.add_run("Date: ")
        r5.bold = True

        r6 = para.add_run(f"{ack_date}")
        r6.bold = False

        for r in [r1, r2, r3, r4, r5, r6]:
            r.font.name = "Arial"
            r.font.size = Pt(7)

    # 👉 SAVE FILE
    output_folder = "media/output"
    os.makedirs(output_folder, exist_ok=True)

    name, ext = os.path.splitext(original_name)
    new_filename = f"{name}_processed{ext}"

    output_path = os.path.join(output_folder, new_filename)
    doc.save(output_path)

    return output_path